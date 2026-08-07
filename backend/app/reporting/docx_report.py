"""Render a stored verdict as a .docx analysis report.

The document is assembled from the same values the API returns and the same
evaluation results the repository publishes. Nothing here is generated prose:
every sentence is either fixed explanatory text about how the system works, or
a number read off the verdict, the corpus, or a recorded evaluation run.

Two structural decisions worth stating. The scope disclaimer sits directly
under the header, before the verdict, because a reader who stops after the
first page must have seen it — a limitations section at the end is where
caveats go to be skipped. And the limitations section reports the classifier's
measured per-class performance rather than a general "models can be wrong"
hedge, since the interesting failure here is specific and one-directional:
SUPPORT recall is 29%, so the support counts in this report are floors.
"""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.config import settings
from app.models import (
    EvidenceItem,
    InsufficientReason,
    Stance,
    Verdict,
    VerdictResponse,
)
from app.reporting import metrics
from app.reporting.corpus_profile import CorpusProfile

# Muted enough to print legibly in greyscale, which a shaded callout has to.
_DISCLAIMER_FILL = "FBE9E7"
_HEADER_FILL = "EDEFF2"
_MUTED = RGBColor(0x55, 0x5F, 0x6D)

_VERDICT_HEADLINE = {
    Verdict.SUPPORTED: "The retrieved evidence supports this claim",
    Verdict.CONTRADICTED: "The retrieved evidence contradicts this claim",
    Verdict.CONFLICTING: "The retrieved literature is split on this claim",
    Verdict.INSUFFICIENT_EVIDENCE: "No verdict was issued",
}

_INSUFFICIENT_HEADLINE = {
    InsufficientReason.NOT_COVERED_BY_CORPUS: "Not covered by this corpus",
    InsufficientReason.EVIDENCE_INCONCLUSIVE: "Evidence inconclusive",
}

_INSUFFICIENT_DETAIL = {
    InsufficientReason.NOT_COVERED_BY_CORPUS: (
        "Nothing in the corpus is about this claim, so no assessment was "
        "attempted. This is a statement about the corpus, not about the claim: "
        "it is not evidence that the claim is false."
    ),
    InsufficientReason.EVIDENCE_INCONCLUSIVE: (
        "The corpus covers this topic, but too little of what was retrieved "
        "took a position either way for a verdict to be issued."
    ),
}

_GRADE_MEANING = {
    "HIGH": "Further evidence is unlikely to change this reading.",
    "MODERATE": "Further evidence could change this reading.",
    "LOW": "Further evidence is likely to change this reading.",
    "VERY_LOW": "Any reading of this evidence is highly uncertain.",
}

_STANCE_LABEL = {
    Stance.SUPPORT: "Supports",
    Stance.CONTRADICT: "Contradicts",
    Stance.NEUTRAL: "Neutral",
}

_TIER_LABEL = {
    "META_ANALYSIS": "Meta-analysis",
    "SYSTEMATIC_REVIEW": "Systematic review",
    "RCT": "Randomised trial",
    "COHORT": "Cohort study",
    "CASE_CONTROL": "Case-control study",
    "CROSS_SECTIONAL": "Cross-sectional study",
    "CASE_REPORT": "Case report",
    "OTHER": "Other design",
}

# Beyond this the pair list stops informing and starts padding; the count is
# reported in full either way.
_MAX_CONFLICT_PAIRS = 25


def build_report(
    response: VerdictResponse,
    corpus: CorpusProfile,
    verified_at: datetime,
    stance_model_source: str,
) -> bytes:
    """Render the verdict as a .docx file and return its bytes."""
    doc = Document()
    _configure_page(doc)
    _set_core_properties(doc, response)

    _write_header(doc, response, verified_at)
    _write_disclaimer(doc)
    _write_verdict(doc, response)
    _write_counts(doc, response)
    _write_explanation(doc, response)
    _write_conflicting_findings(doc, response)
    _write_evidence_table(doc, response)
    _write_limitations(doc, response, corpus)
    _write_methodology(doc, response, corpus, stance_model_source)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def report_filename(response: VerdictResponse) -> str:
    """A filename that says which claim it is without needing to be opened."""
    words = [w for w in "".join(
        c if c.isalnum() or c.isspace() else " " for c in response.claim
    ).split()][:8]
    slug = "-".join(w.lower() for w in words) or "claim"
    date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    return f"evidenceai-{slug}-{date}.docx"


# --- Document scaffolding ---------------------------------------------------


def _configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)


def _set_core_properties(doc: Document, response: VerdictResponse) -> None:
    props = doc.core_properties
    props.title = f"EvidenceAI report — {response.claim}"
    props.author = "EvidenceAI"
    props.comments = (
        "Automated literature triage report. Not clinical, diagnostic, or "
        "treatment advice."
    )


def _shade(element, fill: str) -> None:
    """Apply a background fill to a table cell or paragraph."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def _shade_cell(cell, fill: str) -> None:
    _shade(cell._tc.get_or_add_tcPr(), fill)


def _caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = _MUTED
    run.italic = True


def _kv_table(doc: Document, rows: list[tuple[str, str]], label_width: float = 1.7):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, (label, value) in enumerate(rows):
        left, right = table.rows[index].cells
        left.width = Inches(label_width)
        right.width = Inches(4.9)
        left_run = left.paragraphs[0].add_run(label)
        left_run.bold = True
        left_run.font.size = Pt(9.5)
        right.paragraphs[0].add_run(value).font.size = Pt(9.5)
    return table


# --- Sections ---------------------------------------------------------------


def _write_header(
    doc: Document, response: VerdictResponse, verified_at: datetime
) -> None:
    title = doc.add_heading("Evidence Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    claim = doc.add_paragraph()
    claim_run = claim.add_run(response.claim)
    claim_run.bold = True
    claim_run.font.size = Pt(13)

    _kv_table(
        doc,
        [
            ("Claim assessed", verified_at.strftime("%d %B %Y, %H:%M UTC")),
            (
                "Report generated",
                datetime.now(timezone.utc).strftime("%d %B %Y, %H:%M UTC"),
            ),
            ("System version", f"EvidenceAI {settings.APP_VERSION}"),
            ("Verification ID", response.verification_id),
        ],
    )
    doc.add_paragraph()


def _write_disclaimer(doc: Document) -> None:
    """The scope statement, directly under the header.

    Placed here rather than in the limitations section on purpose: it governs
    how everything below it should be read, and a caveat that arrives after the
    conclusion has already been formed is not doing its job.
    """
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _shade_cell(cell, _DISCLAIMER_FILL)

    heading = cell.paragraphs[0]
    heading_run = heading.add_run("SCOPE OF THIS REPORT")
    heading_run.bold = True
    heading_run.font.size = Pt(10)

    body = cell.add_paragraph()
    body_run = body.add_run(
        "This is a literature triage aid for researchers and students. It is "
        "not a diagnostic tool, and nothing in it is clinical or treatment "
        "advice. It reports what a fixed corpus of published research says "
        "about a claim; it does not decide whether the claim is true, and it "
        "must not be used to make decisions about anyone's care. Every verdict "
        "here describes the retrieved evidence, never a recommended action."
    )
    body_run.font.size = Pt(9.5)
    doc.add_paragraph()


def _write_verdict(doc: Document, response: VerdictResponse) -> None:
    doc.add_heading("Verdict", level=1)

    if response.verdict == Verdict.INSUFFICIENT_EVIDENCE:
        reason = response.insufficient_reason
        headline = (
            _INSUFFICIENT_HEADLINE.get(reason)
            if reason is not None
            else _VERDICT_HEADLINE[Verdict.INSUFFICIENT_EVIDENCE]
        )
        detail = (
            _INSUFFICIENT_DETAIL.get(reason)
            if reason is not None
            else "Too little qualifying evidence was retrieved to issue a verdict."
        )
        _caption(doc, "No verdict issued")
    else:
        headline = _VERDICT_HEADLINE[response.verdict]
        detail = None

    verdict_paragraph = doc.add_paragraph()
    verdict_run = verdict_paragraph.add_run(response.verdict.value.replace("_", " "))
    verdict_run.bold = True
    verdict_run.font.size = Pt(16)

    summary = doc.add_paragraph()
    summary.add_run(headline + ".").font.size = Pt(11)

    if detail:
        doc.add_paragraph(detail)

    grade = response.grade_certainty.value
    rows = [
        (
            "GRADE certainty",
            f"{grade.replace('_', ' ').title()} — {_GRADE_MEANING.get(grade, '')}",
        )
    ]
    if response.coverage is not None:
        rows.append(
            (
                "Corpus coverage",
                f"closest passage {response.coverage.max_cosine:.3f} cosine "
                f"(floor {settings.MIN_MAX_COSINE:.3f}); top "
                f"{response.coverage.top_k} mean "
                f"{response.coverage.mean_topk_cosine:.3f} "
                f"(floor {settings.MIN_MEAN_TOPK_COSINE:.3f})",
            )
        )
    _kv_table(doc, rows)

    _caption(
        doc,
        "Certainty is GRADE-inspired: a deterministic adjustment over the study "
        "designs retrieved, not a formal GRADE assessment by a reviewer.",
    )
    doc.add_paragraph()


def _write_counts(doc: Document, response: VerdictResponse) -> None:
    doc.add_heading("Source counts", level=1)

    paper_count = len({e.paper.paper_id for e in response.evidence})
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"

    headers = [
        "Supporting",
        "Contradicting",
        "Neutral",
        "Passages assessed",
        "Source papers",
    ]
    values = [
        str(response.support_count),
        str(response.contradict_count),
        str(response.neutral_count),
        str(len(response.evidence)),
        str(paper_count),
    ]

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _shade_cell(cell, _HEADER_FILL)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9)

    for index, value in enumerate(values):
        run = table.rows[1].cells[index].paragraphs[0].add_run(value)
        run.font.size = Pt(11)
        run.bold = True

    _caption(
        doc,
        "Counts cover the passages that cleared the relevance threshold. The "
        "verdict is decided on the directional ones only — neutral passages are "
        "reported but never break a tie.",
    )
    doc.add_paragraph()


def _write_explanation(doc: Document, response: VerdictResponse) -> None:
    doc.add_heading("Explanation", level=1)
    doc.add_paragraph(response.explanation)
    _caption(
        doc,
        "This explanation is assembled from the counts and thresholds above and "
        "from sentences quoted verbatim out of the source passages. No language "
        "model writes any part of it, so each statement traces back to a number "
        "or a quotation reproduced elsewhere in this report.",
    )
    doc.add_paragraph()


def _write_conflicting_findings(doc: Document, response: VerdictResponse) -> None:
    doc.add_heading("Conflicting findings", level=1)

    if not response.conflicting_pairs:
        doc.add_paragraph(
            "No direct conflicts were found: the retrieved evidence does not "
            "contain a supporting and a contradicting passage on this claim."
        )
        doc.add_paragraph()
        return

    by_id = {e.passage.passage_id: e for e in response.evidence}
    pairs = response.conflicting_pairs

    doc.add_paragraph(
        f"{response.support_count} supporting and {response.contradict_count} "
        f"contradicting passages were retrieved, giving {len(pairs)} directly "
        f"opposed pairs. They are listed here rather than mixed into the "
        f"evidence table because a disagreement between two sources is a "
        f"finding in its own right — it survives even when one side is in the "
        f"majority."
    )

    shown = pairs[:_MAX_CONFLICT_PAIRS]
    table = doc.add_table(rows=len(shown) + 1, cols=2)
    table.style = "Table Grid"

    for index, header in enumerate(["Supporting passage", "Contradicting passage"]):
        cell = table.rows[0].cells[index]
        _shade_cell(cell, _HEADER_FILL)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9)

    for row_index, (support_id, contradict_id) in enumerate(shown, start=1):
        row = table.rows[row_index]
        for cell_index, passage_id in enumerate((support_id, contradict_id)):
            cell = row.cells[cell_index]
            cell.width = Inches(3.4)
            item = by_id.get(passage_id)
            if item is None:
                cell.paragraphs[0].add_run(passage_id).font.size = Pt(8.5)
                continue

            quote = cell.paragraphs[0]
            quote_run = quote.add_run(_rationale_text(item))
            quote_run.font.size = Pt(9)

            cite = cell.add_paragraph()
            cite_run = cite.add_run(_short_citation(item))
            cite_run.font.size = Pt(8)
            cite_run.font.color.rgb = _MUTED

    if len(pairs) > len(shown):
        _caption(
            doc,
            f"Showing the first {len(shown)} of {len(pairs)} pairs. Every "
            f"supporting passage is opposed to every contradicting one, so the "
            f"pair count grows as their product; the passages themselves are "
            f"listed in full in the evidence table.",
        )
    doc.add_paragraph()


# --- Evidence table ---------------------------------------------------------


def _write_evidence_table(doc: Document, response: VerdictResponse) -> None:
    """The evidence, on its own landscape page.

    Portrait leaves roughly 6.8 inches of text width, which is not enough for a
    passage column beside four metadata columns without the passage wrapping to
    a ribbon. The section break is a formatting concession to the passage text
    being the important column.
    """
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    doc.add_heading("Evidence", level=1)

    if not response.evidence:
        doc.add_paragraph(
            "No passages cleared the relevance threshold, so there is no "
            "evidence to list. Where the corpus does not cover a claim, the "
            "highest-ranked passages are deliberately withheld: they are the "
            "best of an irrelevant set, and presenting them as evidence beside "
            "a non-verdict is the confusion the coverage gate exists to prevent."
        )
        _restore_portrait(doc)
        return

    ordered = sorted(response.evidence, key=lambda e: e.relevance_score, reverse=True)

    table = doc.add_table(rows=len(ordered) + 1, cols=5)
    table.style = "Table Grid"
    widths = [Inches(4.3), Inches(0.9), Inches(0.8), Inches(2.5), Inches(1.1)]
    headers = ["Passage", "Stance", "Relevance", "Source", "Publication tier"]

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = widths[index]
        _shade_cell(cell, _HEADER_FILL)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9)

    for row_index, item in enumerate(ordered, start=1):
        row = table.rows[row_index]
        for index in range(5):
            row.cells[index].width = widths[index]

        _write_passage_cell(row.cells[0], item)

        stance_run = row.cells[1].paragraphs[0].add_run(_STANCE_LABEL[item.stance])
        stance_run.font.size = Pt(9)
        stance_run.bold = True

        score_run = row.cells[2].paragraphs[0].add_run(f"{item.relevance_score:.2f}")
        score_run.font.size = Pt(9)

        _write_source_cell(row.cells[3], item)

        tier = (
            _TIER_LABEL.get(item.paper.publication_tier.value, "Other design")
            if item.paper.publication_tier is not None
            else "Not classified"
        )
        tier_run = row.cells[4].paragraphs[0].add_run(tier)
        tier_run.font.size = Pt(8.5)
        if item.paper.publication_tier is None:
            tier_run.font.color.rgb = _MUTED

    _caption(
        doc,
        "Bold text inside a passage is the sentence the classifier keyed on for "
        "its stance. Relevance is min-max normalised within this result set: it "
        "orders these passages against each other and is not comparable across "
        "different claims.",
    )
    _restore_portrait(doc)


def _restore_portrait(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def _write_passage_cell(cell, item: EvidenceItem) -> None:
    """Write the passage with its rationale sentences bolded in place.

    Rationale sentences are cut from the passage by the classifier, so an exact
    substring match is the right test; anything it cannot locate is appended
    with a note rather than dropped, because silently discarding the model's
    own stated reason would undo the point of quoting it.
    """
    paragraph = cell.paragraphs[0]
    text = item.passage.text
    spans: list[tuple[int, int]] = []
    unmatched: list[str] = []

    for sentence in item.rationale_sentences:
        needle = sentence.strip()
        if not needle:
            continue
        start = text.find(needle)
        if start == -1:
            unmatched.append(needle)
        else:
            spans.append((start, start + len(needle)))

    cursor = 0
    for start, end in sorted(spans):
        if start < cursor:
            continue
        if start > cursor:
            paragraph.add_run(text[cursor:start]).font.size = Pt(9)
        marked = paragraph.add_run(text[start:end])
        marked.font.size = Pt(9)
        marked.bold = True
        cursor = end
    if cursor < len(text):
        paragraph.add_run(text[cursor:]).font.size = Pt(9)

    for sentence in unmatched:
        note = cell.add_paragraph()
        label = note.add_run("Classifier rationale (not locatable verbatim): ")
        label.font.size = Pt(8)
        label.italic = True
        label.font.color.rgb = _MUTED
        body = note.add_run(sentence)
        body.font.size = Pt(8)
        body.font.color.rgb = _MUTED


def _write_source_cell(cell, item: EvidenceItem) -> None:
    paper = item.paper
    title = cell.paragraphs[0]
    title_run = title.add_run(paper.title)
    title_run.font.size = Pt(8.5)
    title_run.bold = True

    meta = cell.add_paragraph()
    meta_run = meta.add_run(_short_citation(item))
    meta_run.font.size = Pt(8)
    meta_run.font.color.rgb = _MUTED

    flags = []
    if paper.is_preprint:
        flags.append("PREPRINT — NOT PEER REVIEWED")
    if paper.openalex_checked and paper.is_retracted:
        flags.append("RETRACTED — DO NOT TREAT AS EVIDENCE")
    if flags:
        flag_paragraph = cell.add_paragraph()
        flag_run = flag_paragraph.add_run(" · ".join(flags))
        flag_run.font.size = Pt(8)
        flag_run.bold = True

    link = cell.add_paragraph()
    link_run = link.add_run(paper.source_url)
    link_run.font.size = Pt(7.5)
    link_run.font.color.rgb = _MUTED


def _short_citation(item: EvidenceItem) -> str:
    paper = item.paper
    authors = paper.authors
    if not authors:
        lead = "Unattributed"
    elif len(authors) == 1:
        lead = authors[0]
    elif len(authors) == 2:
        lead = f"{authors[0]} and {authors[1]}"
    else:
        lead = f"{authors[0]} et al."
    journal = f", {paper.journal}" if paper.journal else ""
    return f"{lead} ({paper.year}){journal}"


def _rationale_text(item: EvidenceItem) -> str:
    if item.rationale_sentences:
        return " ".join(s.strip() for s in item.rationale_sentences if s.strip())
    return item.passage.text


# --- Limitations and methodology --------------------------------------------


def _write_limitations(
    doc: Document, response: VerdictResponse, corpus: CorpusProfile
) -> None:
    doc.add_heading("Limitations", level=1)
    doc.add_paragraph(
        "These apply to every report this system produces, and are stated in "
        "full rather than summarised, because each one bounds what the verdict "
        "above can be taken to mean."
    )

    doc.add_heading("Corpus size and scope", level=2)
    doc.add_paragraph(
        f"The corpus is {corpus.scope_sentence()}. This is a deliberately "
        f"bounded slice of the mental health literature assembled for this "
        f"system — not the literature itself, and not a systematic search. A "
        f"claim can be well supported in published research and still be "
        f"reported here as unsupported simply because the relevant papers were "
        f"never indexed. No result in this report should be read as a statement "
        f"about the state of the evidence at large."
    )
    tier_summary = ", ".join(
        f"{_TIER_LABEL.get(tier, tier)}: {count}"
        for tier, count in corpus.tier_counts.items()
    )
    doc.add_paragraph(
        f"Study designs represented: {tier_summary or 'none classified'}. A "
        f"further {corpus.unclassified_tier_count} papers carry no verified "
        f"design classification, which limits how much weight the certainty "
        f"grade can place on evidence quality. {corpus.preprint_count} papers "
        f"are preprints and have not been peer reviewed."
    )
    doc.add_paragraph(
        f"Retraction status was checked against OpenAlex for "
        f"{corpus.retraction_checked_count} of {corpus.paper_count} papers, of "
        f"which {corpus.retracted_count} are flagged as retracted. Papers that "
        f"were not checked are not thereby known to be clean."
    )

    doc.add_heading("Stance classifier performance", level=2)
    doc.add_paragraph(
        f"Stance is assigned by {metrics.STANCE_EVAL_MODEL}, measured on the "
        f"{metrics.STANCE_EVAL_DATASET}, n={metrics.STANCE_EVAL_N}. Overall "
        f"accuracy is {metrics.STANCE_EVAL_ACCURACY:.1%} and macro F1 is "
        f"{metrics.STANCE_EVAL_MACRO_F1:.3f}."
    )

    table = doc.add_table(rows=len(metrics.STANCE_PER_CLASS) + 1, cols=5)
    table.style = "Table Grid"
    for index, header in enumerate(["Class", "Precision", "Recall", "F1", "Support"]):
        cell = table.rows[0].cells[index]
        _shade_cell(cell, _HEADER_FILL)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9)

    for row_index, row_metrics in enumerate(metrics.STANCE_PER_CLASS, start=1):
        cells = table.rows[row_index].cells
        values = [
            row_metrics.label,
            f"{row_metrics.precision:.3f}",
            f"{row_metrics.recall:.3f}",
            f"{row_metrics.f1:.3f}",
            str(row_metrics.support),
        ]
        for cell_index, value in enumerate(values):
            run = cells[cell_index].paragraphs[0].add_run(value)
            run.font.size = Pt(9)

    doc.add_paragraph(metrics.STANCE_HEADLINE)

    doc.add_heading("Residual gap in the corpus-coverage gate", level=2)
    doc.add_paragraph(metrics.COVERAGE_RESIDUAL_GAP)

    doc.add_heading("Limitations specific to this claim", level=2)
    if response.limitations:
        for limitation in response.limitations:
            doc.add_paragraph(limitation, style="List Bullet")
    else:
        doc.add_paragraph("None recorded for this evidence set.")

    doc.add_heading("Not implemented", level=2)
    doc.add_paragraph(
        "PICO extraction is not implemented. The claim is retrieved against and "
        "assessed exactly as written, so no population, intervention, "
        "comparison or outcome elements were parsed out of it, and indirectness "
        "between the claim's population and a passage's is not detected."
    )
    doc.add_paragraph()


def _write_methodology(
    doc: Document,
    response: VerdictResponse,
    corpus: CorpusProfile,
    stance_model_source: str,
) -> None:
    doc.add_heading("Appendix: methodology", level=1)

    doc.add_heading("Retrieval", level=2)
    doc.add_paragraph(
        f"Passages are ranked by a hybrid of lexical and dense retrieval. BM25 "
        f"(Okapi) scores lexical overlap; {settings.EMBEDDING_MODEL} embeds the "
        f"claim and every passage, and cosine similarity is computed over an "
        f"exact FAISS inner-product index across all {corpus.passage_count} "
        f"passages, so nothing is pruned before scoring. The two score sets are "
        f"each min-max normalised and fused as "
        f"{settings.BM25_WEIGHT:g} × BM25 + {settings.DENSE_WEIGHT:g} × dense, "
        f"and the top {settings.RETRIEVAL_TOP_K} are kept. Passages scoring "
        f"below {settings.MIN_SIMILARITY:g} on that fused score are then "
        f"discarded as insufficiently relevant."
    )

    doc.add_heading("Corpus-coverage gate", level=2)
    doc.add_paragraph(
        f"The fused score above is normalised per query, which pushes the "
        f"top-ranked passage of every claim toward 1.0 — including a claim the "
        f"corpus knows nothing about. It therefore answers 'is this the best of "
        f"what was found?' and can never answer 'is any of this relevant?'. A "
        f"separate gate answers the second question using raw, "
        f"pre-normalisation cosine, which is comparable across queries: a claim "
        f"is treated as covered only if the closest passage reaches "
        f"{settings.MIN_MAX_COSINE:.3f} and the top "
        f"{settings.COVERAGE_TOP_K} average at least "
        f"{settings.MIN_MEAN_TOPK_COSINE:.3f}. Both floors were calibrated "
        f"empirically against on- and off-corpus claim sets. When the gate "
        f"fails, no stance classification is run and no evidence is reported."
    )

    doc.add_heading("Stance classification", level=2)
    doc.add_paragraph(
        f"Each retrieved passage is classified against the claim by "
        f"{stance_model_source} as a natural-language-inference problem: the "
        f"passage is the premise, the claim is the hypothesis, and entailment, "
        f"contradiction and neutral map onto supports, contradicts and neutral. "
        f"The rationale sentence is found by re-scoring each sentence of the "
        f"passage on its own and keeping whichever scores highest for the "
        f"stance already assigned, so it reads out the classifier itself rather "
        f"than a separate similarity heuristic. The classifier reads one "
        f"passage at a time and never the full paper."
    )

    doc.add_heading("How certainty was derived", level=2)
    doc.add_paragraph(
        f"A verdict is issued only when at least "
        f"{settings.MIN_RELEVANT_SOURCES} passages clear the relevance "
        f"threshold and at least one carries a direction. Support and "
        f"contradict shares are taken over directional passages only, so the "
        f"volume of neutral evidence cannot move a verdict. Shares within "
        f"{settings.TIE_MARGIN:g} of each other are reported as conflicting "
        f"rather than resolved in favour of the larger side."
    )
    doc.add_paragraph(
        "The starting certainty is Moderate for a directional verdict and Low "
        "for a conflicting one. It is then adjusted deterministically: up one "
        "level for at least two meta-analyses or systematic reviews alongside "
        "five or more directional passages; down one level if over half the "
        "evidence is low-tier or unclassified; down one level if over 30% is "
        "preprints; and down one level when the evidence base sits exactly at "
        "the minimum source count. No judgement is applied beyond these rules, "
        "and no reviewer assesses risk of bias, imprecision, or publication "
        "bias as formal GRADE requires."
    )

    doc.add_heading("Reproducibility", level=2)
    _kv_table(
        doc,
        [
            ("Verification ID", response.verification_id),
            ("System version", f"EvidenceAI {settings.APP_VERSION}"),
            ("Embedding model", settings.EMBEDDING_MODEL),
            ("Stance model", stance_model_source),
            ("Corpus", f"{corpus.paper_count} papers / {corpus.passage_count} passages"),
            ("Pipeline runtime", f"{response.response_time_ms / 1000:.2f} s"),
        ],
        label_width=1.9,
    )
