"""FastAPI entrypoint for EvidenceAI.

POST /verify runs the real retrieval + stance classification pipeline
(app.pipeline.service), loaded once at startup. GET /verify/{id} and
GET /verify/{id}/report still return hardcoded sample data — there is no
persistence layer for verdicts yet, so a verdict can't actually be looked
back up by id; adding that is out of scope for now.
"""

import os
import sys
import tempfile
from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path

import psutil
from docx import Document
from fastapi import BackgroundTasks, FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

from app.config import settings
from app.models import (
    ClaimRequest,
    EvidenceItem,
    GradeCertainty,
    Paper,
    Passage,
    PICOClaim,
    SourceDatabase,
    Stance,
    Verdict,
    VerdictResponse,
)
from app.pipeline.service import PipelineService, get_service, init_service

_process = psutil.Process(os.getpid())


@asynccontextmanager
async def lifespan(app: FastAPI):
    print("Loading retrieval index and stance classifier...", file=sys.stderr)
    try:
        init_service()
    except FileNotFoundError as exc:
        print(f"\nSTARTUP FAILED: {exc}\n", file=sys.stderr)
        raise
    print("Pipeline ready.", file=sys.stderr)
    yield


app = FastAPI(
    title="EvidenceAI",
    description="Explainable research claim verification for mental health literature.",
    version="0.1.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def _sample_verdict(verification_id: str, claim: str) -> VerdictResponse:
    """Build a hardcoded but schema-valid VerdictResponse for stub endpoints."""
    supporting_paper = Paper(
        paper_id="pmid_23541163",
        title="Randomized Controlled Trial of Mindfulness Meditation for "
        "Generalized Anxiety Disorder",
        authors=["Hoge EA", "Bui E", "Marques L"],
        year=2013,
        abstract="This randomized controlled trial examined the efficacy of "
        "mindfulness-based stress reduction (MBSR) in adults with "
        "generalized anxiety disorder (GAD).",
        source_url="https://pubmed.ncbi.nlm.nih.gov/23541163/",
        source_database=SourceDatabase.PUBMED,
        publication_type="RCT",
        is_preprint=False,
        doi="10.4088/JCP.12m08083",
        retrieved_at=datetime(2026, 7, 15, 9, 30, 0),
    )
    contradicting_paper = Paper(
        paper_id="psyarxiv_7f3kq",
        title="Null Effects of Brief Mindfulness Training on Anxiety in a "
        "Community Sample",
        authors=["Alvarez T", "Nguyen P"],
        year=2024,
        abstract="A preregistered replication found no significant reduction "
        "in trait anxiety following a 4-week brief mindfulness intervention "
        "relative to an active control.",
        source_url="https://psyarxiv.com/7f3kq",
        source_database=SourceDatabase.PSYARXIV,
        publication_type="RCT",
        is_preprint=True,
        doi=None,
        retrieved_at=datetime(2026, 7, 15, 9, 31, 0),
    )
    neutral_paper = Paper(
        paper_id="pmid_31234567",
        title="Prevalence of Generalized Anxiety Disorder in Primary Care "
        "Settings",
        authors=["Kim S", "Patel R"],
        year=2019,
        abstract="This cross-sectional study estimates GAD prevalence among "
        "primary care patients but does not evaluate any intervention.",
        source_url="https://pubmed.ncbi.nlm.nih.gov/31234567/",
        source_database=SourceDatabase.PUBMED,
        publication_type="observational",
        is_preprint=False,
        doi="10.1001/example.31234567",
        retrieved_at=datetime(2026, 7, 15, 9, 32, 0),
    )

    support_passage = Passage(
        passage_id="pmid_23541163_p003",
        paper_id="pmid_23541163",
        text="Participants receiving MBSR showed significantly greater "
        "reductions in anxiety symptoms compared to the stress management "
        "education control group (p < .01).",
        char_start=1820,
        char_end=1978,
        section="Results",
    )
    contradict_passage = Passage(
        passage_id="psyarxiv_7f3kq_p002",
        paper_id="psyarxiv_7f3kq",
        text="No significant between-group difference in trait anxiety was "
        "observed at post-intervention (p = .41).",
        char_start=990,
        char_end=1096,
        section="Results",
    )
    neutral_passage = Passage(
        passage_id="pmid_31234567_p001",
        paper_id="pmid_31234567",
        text="Approximately 8% of primary care patients screened positive for "
        "generalized anxiety disorder.",
        char_start=210,
        char_end=310,
        section="Abstract",
    )

    evidence = [
        EvidenceItem(
            passage=support_passage,
            paper=supporting_paper,
            relevance_score=0.87,
            stance=Stance.SUPPORT,
            stance_confidence=0.92,
            rationale_sentences=[support_passage.text],
            population_match=True,
        ),
        EvidenceItem(
            passage=contradict_passage,
            paper=contradicting_paper,
            relevance_score=0.81,
            stance=Stance.CONTRADICT,
            stance_confidence=0.78,
            rationale_sentences=[contradict_passage.text],
            population_match=True,
        ),
        EvidenceItem(
            passage=neutral_passage,
            paper=neutral_paper,
            relevance_score=0.52,
            stance=Stance.NEUTRAL,
            stance_confidence=0.65,
            rationale_sentences=[neutral_passage.text],
            population_match=None,
        ),
    ]

    return VerdictResponse(
        verification_id=verification_id,
        claim=claim,
        pico=PICOClaim(
            raw_claim=claim,
            population="adults with generalized anxiety disorder",
            intervention="mindfulness-based stress reduction (MBSR)",
            comparison="active or waitlist control",
            outcome="anxiety symptom severity",
        ),
        verdict=Verdict.CONFLICTING,
        grade_certainty=GradeCertainty.MODERATE,
        support_count=1,
        contradict_count=1,
        neutral_count=1,
        evidence=evidence,
        conflicting_pairs=[(support_passage.passage_id, contradict_passage.passage_id)],
        explanation="Retrieved evidence is mixed: 1 study reports a significant "
        "reduction in anxiety symptoms following mindfulness-based "
        "intervention, while 1 preregistered replication reports no "
        "significant effect. A third source provides background prevalence "
        "data only and does not evaluate the intervention.",
        limitations=[
            "Only 3 relevant sources were retrieved; confidence would improve "
            "with a larger evidence base.",
            "One contradicting source is a preprint and has not been peer "
            "reviewed.",
        ],
        response_time_ms=842.0,
    )


@app.get("/health")
def health() -> dict:
    """Liveness check plus real process/model state — check this before
    recording a demo (see README: Demo Mode) rather than assuming the
    service loaded cleanly just because uvicorn didn't print an error."""
    rss_mb = round(_process.memory_info().rss / (1024 * 1024), 1)

    try:
        service: PipelineService = get_service()
        models = {
            "embedding_model": settings.EMBEDDING_MODEL,
            "stance_model_mode": settings.STANCE_MODEL_MODE,
            "stance_model_source": service.stance_classifier.source,
            "stance_model_device": service.stance_classifier.device,
            "passages_indexed": len(service.retrieval.passages),
        }
        pipeline_loaded = True
    except RuntimeError:
        models = None
        pipeline_loaded = False

    return {
        "status": "ok" if pipeline_loaded else "degraded",
        "pipeline_loaded": pipeline_loaded,
        "rss_mb": rss_mb,
        "models": models,
    }


@app.post("/verify", response_model=VerdictResponse)
def verify_claim(request: ClaimRequest) -> VerdictResponse:
    """Retrieve evidence for a claim, classify it, and return a verdict."""
    return get_service().verify(request.claim)


@app.get("/verify/{verification_id}", response_model=VerdictResponse)
def get_verdict(verification_id: str) -> VerdictResponse:
    """Fetch a previously computed verdict by id.

    Stub: returns hardcoded sample data until verdicts are persisted.
    """
    claim = "Mindfulness meditation reduces symptoms of generalized anxiety disorder."
    return _sample_verdict(verification_id, claim)


@app.get("/verify/{verification_id}/report")
def get_report(verification_id: str, background_tasks: BackgroundTasks) -> FileResponse:
    """Export a verdict as a .docx report.

    Stub: generates a minimal placeholder document until app.reporting
    produces real reports from a VerdictResponse.
    """
    verdict = _sample_verdict(
        verification_id,
        "Mindfulness meditation reduces symptoms of generalized anxiety disorder.",
    )

    document = Document()
    document.add_heading("EvidenceAI Report (stub)", level=1)
    document.add_paragraph(f"Verification ID: {verdict.verification_id}")
    document.add_paragraph(f"Claim: {verdict.claim}")
    document.add_paragraph(f"Verdict: {verdict.verdict.value}")

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    document.save(tmp.name)
    tmp.close()

    background_tasks.add_task(Path(tmp.name).unlink, missing_ok=True)
    return FileResponse(
        tmp.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"evidenceai_report_{verification_id}.docx",
    )
